#!/usr/bin/env python3
"""
CV Pilot — AI-powered CV alignment tool for Sashi Kiran Yeggina
----------------------------------------------------------------
Drop a JD → auto-picks best base CV → enhances all sections →
saves with smart tagging for reuse.

Usage:
  python cv_pilot.py                        # Interactive mode
  python cv_pilot.py --jd job.txt           # JD from file
  python cv_pilot.py --list                 # Show saved CV library
  python cv_pilot.py --reuse "AWS Cloud"    # Find reusable CVs for a role
"""

import os, sys, json, uuid, shutil, argparse, textwrap
from datetime import datetime
from pathlib import Path

# ── Dependency check ──────────────────────────────────────────────────────────
MISSING = []
try:
    import anthropic
except ImportError:
    MISSING.append("anthropic")
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    MISSING.append("python-docx")
try:
    from rich.console import Console
    from rich.panel import Panel
    from rich.table import Table
    from rich.prompt import Prompt, Confirm
    from rich.progress import Progress, SpinnerColumn, TextColumn
    from rich import print as rprint
    from rich.markup import escape
except ImportError:
    MISSING.append("rich")

if MISSING:
    print(f"❌  Missing packages: {', '.join(MISSING)}")
    print(f"   Run:  pip install {' '.join(MISSING)}")
    sys.exit(1)

console = Console()

# ── Config ────────────────────────────────────────────────────────────────────
DEFAULT_CV_FOLDER  = "./my_cvs"
DEFAULT_OUTPUT_DIR = "./aligned_cvs"
INDEX_FILE         = "./cv_index.json"
MODEL              = "claude-sonnet-4-6"
MAX_TOKENS         = 8000


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX utilities
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(path: str) -> dict:
    """
    Returns a structured dict:
      { 'full_text': str,
        'paragraphs': [ {'text': str, 'style': str, 'is_bullet': bool} ] }
    """
    doc = Document(path)
    paragraphs = []
    full_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        is_bullet = "List" in style or "Bullet" in style
        paragraphs.append({"text": text, "style": style, "is_bullet": is_bullet})
        full_lines.append(text)

    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text and text not in full_lines:
                        full_lines.append(text)

    return {"full_text": "\n".join(full_lines), "paragraphs": paragraphs}


def apply_enhancements_to_docx(src_path: str, enhancements: dict, out_path: str):
    """
    Opens the source DOCX, applies text replacements, saves to out_path.
    enhancements keys:
      role_title, summary, professional_skills_bullets,
      technical_skills_updates, job_bullets (dict keyed by "Company | Role")
    """
    doc = Document(src_path)

    # Build a flat replacement map: old_paragraph_text → new_text
    replacements = {}

    # ── Role title: find the paragraph that looks like the current job title ──
    # (usually short, bold, in the top 15 paragraphs)
    if "role_title" in enhancements and enhancements["role_title"]:
        new_title = enhancements["role_title"]
        for i, para in enumerate(doc.paragraphs[:20]):
            t = para.text.strip()
            # Heuristic: short line, not a name, not an email, contains "Engineer/Arch/DevOps/SRE"
            if (5 < len(t) < 80 and i > 0 and
                    any(kw in t for kw in ["Engineer","Architect","DevOps","SRE","Cloud","Manager","Analyst","Specialist","AIOps","Platform"])):
                replacements[t] = new_title
                break

    # ── Summary: find the paragraph after a "Summary" heading ────────────────
    if "summary" in enhancements and enhancements["summary"]:
        in_summary = False
        for para in doc.paragraphs:
            t = para.text.strip()
            if t.lower() in ("summary", "professional summary", "profile"):
                in_summary = True
                continue
            if in_summary and len(t) > 60:
                # This is the summary paragraph
                replacements[t] = enhancements["summary"]
                in_summary = False
                break
            if in_summary and t and "Heading" in (para.style.name or ""):
                in_summary = False  # moved to next section without finding it

    # ── Professional Skills bullets ───────────────────────────────────────────
    if "professional_skills_bullets" in enhancements:
        new_bullets = enhancements["professional_skills_bullets"]
        # Collect existing bullet paragraphs from the Professional Skills section
        in_prof_skills = False
        existing_bullets = []
        for para in doc.paragraphs:
            t = para.text.strip()
            style = para.style.name if para.style else ""
            if any(h in t for h in ["Professional Skills", "Core Competencies", "Key Skills"]):
                in_prof_skills = True
                continue
            if in_prof_skills:
                if not t:
                    continue
                if "Heading" in style or (t and t[0].isupper() and len(t) < 40 and ":" not in t
                                          and "List" not in style and t not in ("Technical Skills",)):
                    in_prof_skills = False
                    break
                if "List" in style or style == "Normal":
                    existing_bullets.append(t)

        # Map old bullets → new bullets (pair them up; extras get dropped)
        for i, old_b in enumerate(existing_bullets):
            if i < len(new_bullets):
                replacements[old_b] = new_bullets[i]

    # ── Job bullet points ─────────────────────────────────────────────────────
    if "job_bullets" in enhancements:
        job_bullets_map = enhancements["job_bullets"]  # {"Company | Role": ["bullet", ...]}

        # We iterate the doc and track which company/role section we're in
        current_key = None
        bullet_idx = {}  # key → how many bullets we've replaced so far

        for para in doc.paragraphs:
            t = para.text.strip()
            style = para.style.name if para.style else ""

            # Detect company/role header
            for jk in job_bullets_map:
                parts = [p.strip() for p in jk.split("|")]
                if any(p in t for p in parts if len(p) > 3):
                    current_key = jk
                    bullet_idx[current_key] = 0
                    break

            # Replace bullet if we're inside a tracked section
            if current_key and ("List" in style or "Bullet" in style or "TableParagraph" in style):
                if t and len(t) > 20:
                    bullets = job_bullets_map.get(current_key, [])
                    idx = bullet_idx.get(current_key, 0)
                    if idx < len(bullets):
                        replacements[t] = bullets[idx]
                        bullet_idx[current_key] = idx + 1

    # ── Apply all replacements to doc ─────────────────────────────────────────
    def replace_para_text(para, new_text):
        """Replace a paragraph's text while preserving the first run's formatting."""
        if not para.runs:
            para.text = new_text
            return
        # Save first run formatting
        first_run = para.runs[0]
        bold  = first_run.bold
        size  = first_run.font.size
        color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.type else None
        font_name = first_run.font.name
        # Clear all runs
        for run in para.runs:
            run.text = ""
        # Set text in first run
        first_run.text = new_text
        first_run.bold = bold
        if size:
            first_run.font.size = size
        if color:
            first_run.font.color.rgb = color
        if font_name:
            first_run.font.name = font_name

    for para in doc.paragraphs:
        t = para.text.strip()
        if t in replacements and replacements[t] != t:
            replace_para_text(para, replacements[t])

    # ── Technical skills table ────────────────────────────────────────────────
    if "technical_skills_updates" in enhancements:
        skills_upd = enhancements["technical_skills_updates"]
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) >= 2:
                    category = row.cells[0].text.strip()
                    if category in skills_upd:
                        # Replace second cell content
                        cell = row.cells[1]
                        for para in cell.paragraphs:
                            if para.text.strip():
                                replace_para_text(para, skills_upd[category])
                                break

    doc.save(out_path)


# ══════════════════════════════════════════════════════════════════════════════
#  Claude API calls
# ══════════════════════════════════════════════════════════════════════════════

def get_client() -> "anthropic.Anthropic":
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        # Try loading from .env file in script directory
        env_file = Path(__file__).parent / ".env"
        if env_file.exists():
            for line in env_file.read_text().splitlines():
                if line.startswith("ANTHROPIC_API_KEY="):
                    api_key = line.split("=", 1)[1].strip().strip('"').strip("'")
                    break
    if not api_key:
        console.print(Panel(
            "[bold red]ANTHROPIC_API_KEY not found.[/]\n\n"
            "Set it in a [bold].env[/] file next to cv_pilot.py:\n"
            "  [cyan]ANTHROPIC_API_KEY=sk-ant-...[/]\n\n"
            "Or export it in your terminal:\n"
            "  [cyan]export ANTHROPIC_API_KEY=sk-ant-...[/]",
            title="❌  API Key Missing", border_style="red"
        ))
        sys.exit(1)
    return anthropic.Anthropic(api_key=api_key)


def call_claude(client, system: str, user: str, json_mode: bool = True) -> str:
    """Call Claude and return the response text."""
    msgs = [{"role": "user", "content": user}]
    response = client.messages.create(
        model=MODEL,
        max_tokens=MAX_TOKENS,
        system=system,
        messages=msgs,
    )
    return response.content[0].text


def parse_jd(client, jd_text: str) -> dict:
    system = textwrap.dedent("""\
        You are a senior technical recruiter with 15 years of experience placing
        DevOps, Cloud, and SRE engineers. Analyze job descriptions with precision.
    """)
    user = textwrap.dedent(f"""\
        Analyze this job description and return a JSON object with EXACTLY these keys:

        {{
          "role_title": "clean job title, e.g. AWS Cloud Engineer",
          "role_category": "one of: AWS Cloud Engineer | Cloud Solutions Architect | DevOps Engineer | SRE | AIOps Engineer | Azure DevOps Engineer | Cloud Network Engineer | Platform Engineer",
          "seniority": "Junior | Mid | Senior | Principal | Staff",
          "required_skills": ["skill1", "skill2", ...],
          "preferred_skills": ["skill1", ...],
          "years_experience": "e.g. 5-8 years",
          "key_responsibilities": ["resp1", "resp2", ...],
          "ats_keywords": ["keyword1", ...],
          "company_type": "startup | scale-up | enterprise | unknown",
          "cloud_platform": "AWS | Azure | GCP | Multi-cloud | unknown"
        }}

        JOB DESCRIPTION:
        {jd_text}

        Return ONLY valid JSON. No markdown, no explanation.
    """)
    raw = call_claude(client, system, user)
    # Strip markdown fences if present
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


def score_cv_against_jd(client, cv_summaries: list, jd_analysis: dict) -> list:
    """
    cv_summaries: [{"filename": str, "text_excerpt": str}, ...]
    Returns sorted list: [{"filename": str, "score": int, "reasoning": str, "key_matches": [], "key_gaps": []}, ...]
    """
    summaries_block = "\n\n".join([
        f"CV #{i+1} — {cv['filename']}:\n{cv['text_excerpt'][:1500]}"
        for i, cv in enumerate(cv_summaries)
    ])
    system = "You are an expert technical recruiter who evaluates CV-to-JD fit with precision and speed."
    user = textwrap.dedent(f"""\
        Score each of these CVs against the job requirements below. Return a JSON array.

        JOB REQUIREMENTS:
        - Role: {jd_analysis['role_title']} ({jd_analysis['seniority']})
        - Cloud platform: {jd_analysis['cloud_platform']}
        - Required skills: {', '.join(jd_analysis['required_skills'])}
        - Preferred skills: {', '.join(jd_analysis.get('preferred_skills', []))}
        - ATS keywords: {', '.join(jd_analysis['ats_keywords'])}

        CVs TO SCORE:
        {summaries_block}

        Return a JSON array sorted by score descending, each item:
        {{
          "filename": "...",
          "score": <integer 0-100>,
          "reasoning": "2-3 sentences why this CV is or isn't a good match",
          "key_matches": ["skill1", ...],
          "key_gaps": ["missing1", ...]
        }}

        Return ONLY valid JSON array. No markdown.
    """)
    raw = call_claude(client, system, user)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    scored = json.loads(raw.strip())
    return sorted(scored, key=lambda x: x["score"], reverse=True)


def enhance_cv(client, cv_data: dict, jd_analysis: dict) -> dict:
    """
    cv_data: {"full_text": str, "paragraphs": [...]}
    Returns enhancements dict ready for apply_enhancements_to_docx()
    """
    cv_text = cv_data["full_text"]
    system = textwrap.dedent("""\
        You are an elite technical CV writer who specialises in DevOps, Cloud, and SRE engineering CVs.
        You write in a natural, human voice — not robotic or formulaic.
        You NEVER fabricate experience. You only reframe, reorder, and emphasise existing content.
        Your goal: make this engineer's real experience land perfectly for the target role.
    """)
    user = textwrap.dedent(f"""\
        Enhance this engineer's CV to align with the target job. Follow all rules strictly.

        RULES:
        1. NEVER invent skills, tools, or experience that aren't in the original CV
        2. Use natural, first-person-implied language — as if the engineer wrote it themselves
        3. Mirror JD keywords naturally — don't stuff them mechanically
        4. Make bullet points sound like real project stories, not job descriptions
        5. Summary: 3-5 sentences, confident tone, mentions 2-3 specific JD requirements the engineer actually has
        6. Professional skills: reorder to put JD-relevant skills first; drop unrelated skills; max 30 bullets
        7. Job bullets: rewrite to emphasise JD-relevant accomplishments; keep authentic voice; 4-8 bullets per role
        8. Role title: update to match the JD title exactly (only if the engineer genuinely qualifies)

        TARGET JOB:
        - Title: {jd_analysis['role_title']}
        - Cloud: {jd_analysis['cloud_platform']}
        - Required skills: {', '.join(jd_analysis['required_skills'])}
        - ATS keywords to naturally include: {', '.join(jd_analysis['ats_keywords'])}
        - Key responsibilities: {chr(10).join('  - ' + r for r in jd_analysis['key_responsibilities'][:6])}

        ORIGINAL CV:
        {cv_text[:6000]}

        Return a JSON object with EXACTLY these keys:
        {{
          "role_title": "New role title string",
          "summary": "Full rewritten summary paragraph",
          "professional_skills_bullets": [
            "13+ years ...",
            "Strong experience ...",
            ...
          ],
          "technical_skills_updates": {{
            "Cloud Platforms": "updated skills string for this row",
            "Containerization & Orchestration": "...",
            "CI/CD & SDLC": "...",
            "Monitoring & Observability": "..."
          }},
          "job_bullets": {{
            "Amazon": [
              "Bullet point 1 rewritten naturally",
              "Bullet point 2 ...",
              ...
            ],
            "AT&T": ["...", "..."],
            "Airbus": ["...", "..."],
            "Bank of America": ["...", "..."]
          }}
        }}

        Return ONLY valid JSON. No markdown fences. No explanation.
    """)
    raw = call_claude(client, system, user)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip())


# ══════════════════════════════════════════════════════════════════════════════
#  CV Index (library / reuse store)
# ══════════════════════════════════════════════════════════════════════════════

def load_index(index_path: str) -> dict:
    if os.path.exists(index_path):
        with open(index_path) as f:
            return json.load(f)
    return {"cvs": []}


def save_index(index: dict, index_path: str):
    with open(index_path, "w") as f:
        json.dump(index, f, indent=2)


def add_to_index(index: dict, entry: dict, index_path: str):
    index["cvs"].append(entry)
    save_index(index, index_path)


def find_reusable_cv(client, index: dict, jd_analysis: dict) -> dict | None:
    """
    Check if any saved CV in the library is a close enough match (≥75)
    to be reused without rebuilding from scratch.
    """
    if not index["cvs"]:
        return None

    # Build a quick summary of saved CVs for Claude to evaluate
    summaries = []
    for entry in index["cvs"]:
        summaries.append({
            "id": entry["id"],
            "filename": entry["filename"],
            "role_type": entry["role_type"],
            "jd_keywords": entry.get("jd_keywords", []),
            "reuse_tags": entry.get("reuse_tags", []),
            "ats_score": entry.get("ats_score", 0),
        })

    system = "You are a smart CV librarian. You match saved CVs to new job requirements."
    user = textwrap.dedent(f"""\
        I have a library of saved, aligned CVs. For this new job, tell me if any existing CV
        is close enough to reuse (score ≥ 75) instead of building from scratch.

        NEW JOB:
        - Role: {jd_analysis['role_title']}
        - Cloud: {jd_analysis['cloud_platform']}
        - Required skills: {', '.join(jd_analysis['required_skills'][:10])}

        SAVED CVs:
        {json.dumps(summaries, indent=2)}

        Return JSON:
        {{
          "best_match_id": "uuid or null",
          "match_score": <int 0-100>,
          "reasoning": "short explanation"
        }}

        Return ONLY valid JSON.
    """)
    raw = call_claude(client, system, user)
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    result = json.loads(raw.strip())

    if result.get("match_score", 0) >= 75 and result.get("best_match_id"):
        for entry in index["cvs"]:
            if entry["id"] == result["best_match_id"]:
                return {**entry, "match_score": result["match_score"], "reasoning": result["reasoning"]}
    return None


def build_smart_filename(jd_analysis: dict, company: str = "") -> str:
    """Generate a structured filename for the saved CV."""
    role = jd_analysis["role_title"].replace(" ", "_").replace("/", "-")
    platform = jd_analysis["cloud_platform"].replace("-", "")
    date = datetime.now().strftime("%Y%m%d")
    company_part = f"_{company.replace(' ', '_')}" if company else ""
    return f"{role}_{platform}{company_part}_{date}.docx"


# ══════════════════════════════════════════════════════════════════════════════
#  Display helpers
# ══════════════════════════════════════════════════════════════════════════════

def show_banner():
    console.print(Panel(
        "[bold cyan]CV Pilot[/]  [dim]— AI-powered CV alignment engine[/]\n"
        "[dim]Powered by Claude AI  |  github.com/sashi[/]",
        border_style="cyan", padding=(0, 2)
    ))


def show_jd_summary(jd: dict):
    t = Table(title="📋  JD Analysis", show_header=False, border_style="blue", padding=(0, 1))
    t.add_column("Key", style="bold cyan", width=22)
    t.add_column("Value")
    t.add_row("Role", jd["role_title"])
    t.add_row("Category", jd["role_category"])
    t.add_row("Seniority", jd["seniority"])
    t.add_row("Cloud Platform", jd["cloud_platform"])
    t.add_row("Experience", jd.get("years_experience", "Not specified"))
    t.add_row("Required Skills", ", ".join(jd["required_skills"][:8]))
    t.add_row("ATS Keywords", ", ".join(jd["ats_keywords"][:10]))
    console.print(t)


def show_scores(scored: list):
    t = Table(title="🎯  CV Match Scores", border_style="green", padding=(0, 1))
    t.add_column("#", width=4, justify="right")
    t.add_column("CV File", style="bold")
    t.add_column("Score", justify="center", width=8)
    t.add_column("Key Matches", width=35)
    t.add_column("Key Gaps", width=30)
    for i, row in enumerate(scored):
        score = row["score"]
        color = "green" if score >= 75 else "yellow" if score >= 55 else "red"
        t.add_row(
            str(i + 1),
            Path(row["filename"]).stem,
            f"[{color}]{score}/100[/]",
            ", ".join(row.get("key_matches", [])[:4]),
            ", ".join(row.get("key_gaps", [])[:3]),
        )
    console.print(t)


def show_library(index: dict):
    if not index["cvs"]:
        console.print("[yellow]No CVs saved yet. Run cv_pilot.py to align and save your first CV.[/]")
        return
    t = Table(title=f"📚  CV Library  ({len(index['cvs'])} saved)", border_style="magenta", padding=(0, 1))
    t.add_column("Date", width=12)
    t.add_column("Filename", style="bold")
    t.add_column("Role Type", width=28)
    t.add_column("ATS Score", justify="center", width=10)
    t.add_column("Base CV", width=20)
    t.add_column("Tags", width=35)
    for entry in sorted(index["cvs"], key=lambda x: x.get("created_date",""), reverse=True):
        score = entry.get("ats_score", "—")
        color = "green" if isinstance(score, int) and score >= 80 else "yellow"
        t.add_row(
            entry.get("created_date", "—")[:10],
            Path(entry["filename"]).name[:40],
            entry.get("role_type", "—"),
            f"[{color}]{score}[/]" if score != "—" else "—",
            Path(entry.get("base_cv", "—")).stem[:20],
            ", ".join(entry.get("reuse_tags", [])[:5]),
        )
    console.print(t)


# ══════════════════════════════════════════════════════════════════════════════
#  Main workflow
# ══════════════════════════════════════════════════════════════════════════════

def run_alignment(jd_text: str, cv_folder: str, output_dir: str, index_path: str,
                  company: str = "", auto: bool = False):
    """Core pipeline: JD → best CV → enhance → save."""
    client = get_client()
    index  = load_index(index_path)
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # ── Step 1: Parse the JD ─────────────────────────────────────────────────
    with Progress(SpinnerColumn(), TextColumn("[cyan]{task.description}"), transient=True) as p:
        p.add_task("Analysing job description...", total=None)
        jd_analysis = parse_jd(client, jd_text)

    console.print()
    show_jd_summary(jd_analysis)
    console.print()

    # ── Step 2: Check reuse library ───────────────────────────────────────────
    reuse = find_reusable_cv(client, index, jd_analysis)
    if reuse:
        console.print(Panel(
            f"[bold green]♻️  Found a reusable CV![/]\n\n"
            f"  [bold]{reuse['filename']}[/]\n"
            f"  Match score: [green]{reuse['match_score']}/100[/]\n"
            f"  {reuse['reasoning']}",
            title="Reuse Library Hit", border_style="green"
        ))
        if not auto:
            use_existing = Confirm.ask("Use this existing CV as the base (saves time)?", default=True)
        else:
            use_existing = True

        if use_existing:
            console.print(f"\n[green]✓  Using existing: {reuse['filename']}[/]")
            existing_path = os.path.join(output_dir, reuse["filename"])
            if os.path.exists(existing_path):
                console.print("[dim]Minor tune-up of existing CV for this specific JD...[/]")
                src_for_enhancement = existing_path
            else:
                src_for_enhancement = None
                console.print("[yellow]File not found locally; will rebuild from base CV.[/]")
        else:
            src_for_enhancement = None
    else:
        src_for_enhancement = None

    # ── Step 3: Score CVs in the folder ──────────────────────────────────────
    cv_files = list(Path(cv_folder).glob("*.docx")) + list(Path(cv_folder).glob("*.DOCX"))
    if not cv_files:
        console.print(f"[red]No DOCX files found in {cv_folder}[/]")
        sys.exit(1)

    console.print(f"[dim]Found {len(cv_files)} CVs in {cv_folder}[/]")

    cv_summaries = []
    with Progress(SpinnerColumn(), TextColumn("[cyan]{task.description}"), transient=True) as p:
        task = p.add_task(f"Reading {len(cv_files)} CVs...", total=None)
        for f in cv_files:
            try:
                data = extract_text_from_docx(str(f))
                cv_summaries.append({
                    "filename": str(f),
                    "text_excerpt": data["full_text"][:2000],
                })
            except Exception as e:
                console.print(f"[yellow]⚠  Could not read {f.name}: {e}[/]")

    with Progress(SpinnerColumn(), TextColumn("[cyan]{task.description}"), transient=True) as p:
        p.add_task("Scoring CVs against JD...", total=None)
        scored = score_cv_against_jd(client, cv_summaries, jd_analysis)

    show_scores(scored)
    console.print()

    # ── Step 4: Pick best CV ─────────────────────────────────────────────────
    if auto:
        chosen_path = scored[0]["filename"]
    else:
        best = scored[0]
        console.print(f"[bold green]Best match:[/] {Path(best['filename']).name}  "
                      f"[green]({best['score']}/100)[/]")
        console.print(f"[dim]{best['reasoning']}[/]\n")
        if Confirm.ask("Use this CV as the base?", default=True):
            chosen_path = best["filename"]
        else:
            idx = int(Prompt.ask("Enter the # of the CV to use instead", default="1")) - 1
            chosen_path = scored[min(idx, len(scored)-1)]["filename"]

    if src_for_enhancement is None:
        src_for_enhancement = chosen_path

    console.print(f"\n[cyan]Base CV:[/] {Path(chosen_path).name}\n")

    # ── Step 5: Extract and enhance ───────────────────────────────────────────
    cv_data = extract_text_from_docx(src_for_enhancement)

    with Progress(SpinnerColumn(), TextColumn("[cyan]{task.description}"), transient=True) as p:
        p.add_task("Enhancing CV with Claude AI...", total=None)
        enhancements = enhance_cv(client, cv_data, jd_analysis)

    console.print("[green]✓  Enhancement complete[/]")
    console.print(f"   Role title → [bold]{enhancements.get('role_title', '—')}[/]")
    console.print(f"   Summary    → rewritten ({len(enhancements.get('summary',''))} chars)")
    console.print(f"   Skills     → {len(enhancements.get('professional_skills_bullets',[]))} bullets updated")
    job_b = enhancements.get("job_bullets", {})
    console.print(f"   Job bullets→ {sum(len(v) for v in job_b.values())} points across {len(job_b)} roles")
    console.print()

    # ── Step 6: Apply and save ────────────────────────────────────────────────
    filename  = build_smart_filename(jd_analysis, company)
    out_path  = os.path.join(output_dir, filename)

    apply_enhancements_to_docx(src_for_enhancement, enhancements, out_path)

    # ── Step 7: Update index ──────────────────────────────────────────────────
    entry = {
        "id": str(uuid.uuid4()),
        "filename": filename,
        "base_cv": Path(chosen_path).name,
        "role_type": jd_analysis["role_title"],
        "role_category": jd_analysis["role_category"],
        "cloud_platform": jd_analysis["cloud_platform"],
        "company": company,
        "created_date": datetime.now().isoformat(),
        "jd_keywords": jd_analysis["ats_keywords"],
        "required_skills": jd_analysis["required_skills"],
        "ats_score": scored[0]["score"],
        "reuse_tags": list(set(
            [jd_analysis["cloud_platform"].lower()]
            + [s.lower() for s in jd_analysis["required_skills"][:8]]
            + [jd_analysis["role_category"].lower().replace(" ", "_")]
        )),
    }
    add_to_index(index, entry, index_path)

    # ── Step 8: Auto-sync to GitHub (if configured) ───────────────────────────
    github_line = ""
    try:
        from cv_sync import CVSync
        sync = CVSync(script_dir=str(Path(__file__).parent))
        if sync.enabled:
            with Progress(SpinnerColumn(), TextColumn("[cyan]{task.description}"), transient=True) as p:
                p.add_task("Syncing to GitHub...", total=None)
                sync.push_cv(out_path, role_category=jd_analysis.get("role_category", ""), date=entry["created_date"][:10])
                sync.push_index(index_path)
            github_line = f"\n  ☁   Synced to GitHub: [cyan]{sync.repo}[/]"
    except Exception:
        pass  # GitHub sync is optional; never fail the main flow

    console.print(Panel(
        f"[bold green]✅  CV saved successfully![/]\n\n"
        f"  📄  [bold]{filename}[/]\n"
        f"  📁  {output_dir}\n"
        f"{github_line}\n"
        f"  Added to library index: [cyan]{index_path}[/]\n"
        f"  ({len(index['cvs'])} CVs total in library)",
        title="Done", border_style="green"
    ))
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
#  CLI entry point
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="CV Pilot — AI CV alignment engine",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            Examples:
              python cv_pilot.py                         # interactive mode
              python cv_pilot.py --jd amazon_jd.txt     # JD from file
              python cv_pilot.py --jd amazon_jd.txt --company Amazon --auto
              python cv_pilot.py --list                  # show CV library
              python cv_pilot.py --reuse "AWS Cloud"     # find reusable CVs
        """)
    )
    parser.add_argument("--jd",       type=str, help="Path to JD text file")
    parser.add_argument("--company",  type=str, default="", help="Target company name (optional)")
    parser.add_argument("--cv-folder",type=str, default=DEFAULT_CV_FOLDER, help="Folder containing your base CVs")
    parser.add_argument("--output",   type=str, default=DEFAULT_OUTPUT_DIR, help="Output folder for aligned CVs")
    parser.add_argument("--index",    type=str, default=INDEX_FILE, help="Path to CV index JSON file")
    parser.add_argument("--list",     action="store_true", help="List all CVs in library and exit")
    parser.add_argument("--reuse",    type=str, help="Show library CVs matching a role keyword")
    parser.add_argument("--auto",     action="store_true", help="Auto-select best CV without prompts")
    args = parser.parse_args()

    show_banner()
    console.print()

    # ── --list mode ───────────────────────────────────────────────────────────
    if args.list:
        index = load_index(args.index)
        show_library(index)
        return

    # ── --reuse search mode ───────────────────────────────────────────────────
    if args.reuse:
        index = load_index(args.index)
        keyword = args.reuse.lower()
        matches = [e for e in index["cvs"]
                   if keyword in e.get("role_type", "").lower()
                   or keyword in " ".join(e.get("reuse_tags", [])).lower()]
        if matches:
            filtered = {"cvs": matches}
            show_library(filtered)
        else:
            console.print(f"[yellow]No CVs found matching '{args.reuse}'[/]")
        return

    # ── Alignment mode ────────────────────────────────────────────────────────
    if not os.path.isdir(args.cv_folder):
        console.print(Panel(
            f"[red]CV folder not found:[/] {args.cv_folder}\n\n"
            f"Create it and put your base DOCX CVs inside:\n"
            f"  [cyan]mkdir {args.cv_folder}[/]\n"
            f"  (then copy your 8+ CVs there)",
            title="❌  Setup needed", border_style="red"
        ))
        sys.exit(1)

    # Get JD text
    if args.jd:
        if not os.path.exists(args.jd):
            console.print(f"[red]JD file not found:[/] {args.jd}")
            sys.exit(1)
        jd_text = Path(args.jd).read_text(encoding="utf-8")
        console.print(f"[dim]JD loaded from {args.jd} ({len(jd_text)} chars)[/]\n")
    else:
        console.print("[bold cyan]Paste the job description below.[/]")
        console.print("[dim]Press Enter twice when done, then type END and press Enter:[/]\n")
        lines = []
        while True:
            try:
                line = input()
            except EOFError:
                break
            if line.strip().upper() == "END":
                break
            lines.append(line)
        jd_text = "\n".join(lines).strip()
        if not jd_text:
            console.print("[red]No JD text provided. Exiting.[/]")
            sys.exit(1)

    company = args.company
    if not company and not args.auto:
        company = Prompt.ask("Target company name? (press Enter to skip)", default="")

    run_alignment(
        jd_text=jd_text,
        cv_folder=args.cv_folder,
        output_dir=args.output,
        index_path=args.index,
        company=company,
        auto=args.auto,
    )


if __name__ == "__main__":
    main()
