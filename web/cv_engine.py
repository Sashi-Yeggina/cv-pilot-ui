"""
cv_engine.py — Claude API logic for CV Pilot Web UI
----------------------------------------------------
Functions:
  parse_jd(client, jd_text)          -> jd analysis dict
  score_cvs(client, cv_list, jd)     -> sorted scored list
  enhance_cv(client, cv_text, jd)    -> enhancements dict
  apply_enhancements(cv_bytes, enhancements) -> bytes (DOCX)
"""

import json
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── JD Parsing ──────────────────────────────────────────────────────────────

def parse_jd(client, jd_text: str) -> dict:
    """
    Parse a job description and extract structured info.
    Returns dict with role metadata, skills, keywords, etc.
    """
    prompt = f"""Analyse this job description and return structured JSON.

Job Description:
{jd_text}

Return ONLY valid JSON with this exact structure:
{{
  "role_title": "exact role title from JD",
  "role_category": "one of: AWS Cloud Engineer | Cloud Solutions Architect | DevOps Engineer | SRE | AIOps Engineer | Azure DevOps Engineer | Cloud Network Engineer | Platform Engineer | Other",
  "seniority": "one of: Junior | Mid | Senior | Principal | Lead",
  "required_skills": ["skill1", "skill2", ...],
  "preferred_skills": ["skill1", "skill2", ...],
  "years_experience": "e.g. 5+ years",
  "key_responsibilities": ["resp1", "resp2", ...],
  "ats_keywords": ["keyword1", "keyword2", ...],
  "company_type": "e.g. Enterprise | Startup | Consultancy | Bank | Telco",
  "cloud_platform": "one of: AWS | Azure | GCP | MultiCloud | On-Prem | Hybrid"
}}"""

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )

    text = response.content[0].text.strip()
    # Strip markdown code fences if present
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
    return json.loads(text.strip())


# ─── CV Scoring ───────────────────────────────────────────────────────────────

def score_cvs(client, cv_list: list, jd: dict) -> list:
    """
    Score each CV against the JD.
    cv_list: list of dicts with keys: id, filename, text
    Returns sorted list (highest score first) with score, reasoning, matches, gaps added.
    """
    if not cv_list:
        return []

    # Build a compact summary of each CV for scoring
    cv_summaries = []
    for i, cv in enumerate(cv_list):
        # Truncate text to avoid token limits — first 3000 chars is plenty for scoring
        snippet = cv["text"][:3000]
        cv_summaries.append(f"CV_{i}: {cv['filename']}\n{snippet}")

    joined = "\n\n---\n\n".join(cv_summaries)

    jd_summary = (
        f"Role: {jd.get('role_title', '')}\n"
        f"Platform: {jd.get('cloud_platform', '')}\n"
        f"Required Skills: {', '.join(jd.get('required_skills', []))}\n"
        f"Preferred Skills: {', '.join(jd.get('preferred_skills', []))}\n"
        f"ATS Keywords: {', '.join(jd.get('ats_keywords', []))}\n"
        f"Responsibilities: {'; '.join(jd.get('key_responsibilities', []))}"
    )

    prompt = f"""Score each CV against this job description. Return ONLY valid JSON.

JD Summary:
{jd_summary}

CVs to score:
{joined}

Return a JSON array sorted by score descending:
[
  {{
    "cv_index": 0,
    "score": 85,
    "reasoning": "Strong AWS and Kubernetes experience matching JD requirements",
    "key_matches": ["Kubernetes", "AWS EKS", "Terraform"],
    "key_gaps": ["Service Mesh experience", "Go programming"]
  }},
  ...
]

Scoring rubric:
- 90-100: Near perfect match, most required + preferred skills present
- 75-89: Strong match, most required skills present
- 60-74: Good match, some required skills present
- 45-59: Partial match, limited overlap
- Below 45: Poor match

Return ONLY the JSON array, no other text."""

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    text = response.content[0].text.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    scored = json.loads(text.strip())

    # Merge scores back into cv_list
    result = []
    for item in scored:
        idx = item["cv_index"]
        if idx < len(cv_list):
            merged = dict(cv_list[idx])
            merged["score"] = item.get("score", 0)
            merged["reasoning"] = item.get("reasoning", "")
            merged["key_matches"] = item.get("key_matches", [])
            merged["key_gaps"] = item.get("key_gaps", [])
            result.append(merged)

    return sorted(result, key=lambda x: x.get("score", 0), reverse=True)


# ─── CV Enhancement ──────────────────────────────────────────────────────────

def enhance_cv(client, cv_text: str, jd: dict) -> dict:
    """
    Rewrite CV sections to align with the JD.
    Returns enhancements dict with updated sections.
    """
    jd_str = json.dumps(jd, indent=2)

    prompt = f"""You are a professional CV writer. Rewrite specific sections of this CV to align with the job description.

RULES:
- NEVER fabricate experience or skills the candidate doesn't have
- Reframe existing experience to highlight JD-relevant aspects
- Mirror JD keywords naturally (ATS optimised) — don't keyword-stuff
- Use natural, human tone — avoid robotic phrases like "Responsible for..." or "Involved in..."
- Start bullets with strong action verbs: Led, Built, Designed, Automated, Reduced, Improved...
- Keep all dates, company names, job titles factually accurate
- Professional summary: 3-4 sentences, confident but not arrogant

JOB DESCRIPTION ANALYSIS:
{jd_str}

CANDIDATE'S CURRENT CV:
{cv_text[:6000]}

Return ONLY valid JSON with this structure:
{{
  "role_title": "Updated role title matching the JD (e.g. Senior DevOps Engineer)",
  "summary": "3-4 sentence professional summary tailored to this JD",
  "professional_skills_bullets": [
    "Cloud Platforms: AWS (EKS, EC2, RDS, Lambda, S3), Terraform, CloudFormation",
    "CI/CD: Jenkins, GitHub Actions, ArgoCD, GitLab CI",
    "Containers: Docker, Kubernetes, Helm",
    "Monitoring: Prometheus, Grafana, Datadog, ELK Stack"
  ],
  "technical_skills_updates": {{
    "add": ["skill1", "skill2"],
    "emphasise": ["skill3", "skill4"]
  }},
  "job_bullets": {{
    "Company Name 1": [
      "Action-verb led bullet point that mirrors JD keywords naturally...",
      "Another strong bullet point..."
    ],
    "Company Name 2": [
      "Bullet point..."
    ]
  }}
}}

For job_bullets: use the EXACT company names as they appear in the CV. Only include companies where you can genuinely improve the bullets based on the JD. Provide 3-5 bullets per company."""

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )

    text = response.content[0].text.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    return json.loads(text.strip())


# ─── Apply Enhancements to DOCX bytes ────────────────────────────────────────

def apply_enhancements(cv_bytes: bytes, enhancements: dict) -> bytes:
    """
    Apply enhancements dict to a DOCX (provided as bytes).
    Returns updated DOCX as bytes.

    Strategy: Replace matching paragraphs using heuristics:
    - Role title: short bold line near top with engineering keywords
    - Summary: paragraph(s) after a "Summary" or "Profile" heading
    - Job bullets: paragraphs starting with "•" or list style, near company name
    """
    doc = Document(io.BytesIO(cv_bytes))
    paras = doc.paragraphs

    role_title = enhancements.get("role_title", "")
    summary = enhancements.get("summary", "")
    job_bullets = enhancements.get("job_bullets", {})
    skills_bullets = enhancements.get("professional_skills_bullets", [])

    # ── Pass 1: build a map of paragraph text → index ──────────────────────
    para_texts = [p.text.strip() for p in paras]

    # ── Helper: replace paragraph text preserving first run formatting ──────
    def replace_para_text(para, new_text: str):
        if not para.runs:
            para.text = new_text
            return
        # Preserve formatting of first run
        first_run = para.runs[0]
        fmt = {
            "bold": first_run.bold,
            "italic": first_run.italic,
            "font_size": first_run.font.size,
            "font_name": first_run.font.name,
            "color": first_run.font.color.rgb if first_run.font.color and first_run.font.color.type else None,
        }
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_text
        para.runs[0].bold = fmt["bold"]
        para.runs[0].italic = fmt["italic"]
        if fmt["font_size"]:
            para.runs[0].font.size = fmt["font_size"]
        if fmt["font_name"]:
            para.runs[0].font.name = fmt["font_name"]

    # ── Role title replacement ───────────────────────────────────────────────
    if role_title:
        title_keywords = ["engineer", "architect", "devops", "sre", "cloud",
                          "platform", "consultant", "developer", "manager",
                          "specialist", "analyst", "lead", "director"]
        # Skip paragraphs that contain email/phone — those are contact lines
        contact_indicators = ["@", "tel:", "phone:", "+1", "linkedin", "github"]
        for i, para in enumerate(paras[:25]):
            text = para.text.strip()
            text_lower = text.lower()
            has_contact = any(ci in text_lower for ci in contact_indicators)
            # Allow up to 150 chars, must have title keyword, must not be a contact line
            if (0 < len(text) < 150 and
                    any(kw in text_lower for kw in title_keywords) and
                    not has_contact and
                    "|" in text):   # role title lines typically use | separator
                replace_para_text(para, role_title)
                break
        else:
            # Fallback: first short line with title keyword, no email
            for i, para in enumerate(paras[:25]):
                text = para.text.strip()
                text_lower = text.lower()
                has_contact = any(ci in text_lower for ci in contact_indicators)
                if (0 < len(text) < 120 and
                        any(kw in text_lower for kw in title_keywords) and
                        not has_contact):
                    replace_para_text(para, role_title)
                    break

    # ── Summary replacement ──────────────────────────────────────────────────
    if summary:
        summary_keywords = ["summary", "profile", "objective", "about me",
                            "professional summary", "executive summary"]
        summary_replaced = False

        # Strategy A: find a paragraph whose ENTIRE text is a summary heading
        # (short heading-only para like "Summary" or "Professional Summary")
        for i, para in enumerate(paras[:40]):
            text_lower = para.text.lower().strip()
            is_pure_heading = any(text_lower == kw for kw in summary_keywords)
            if is_pure_heading:
                # Replace the next non-empty paragraph (which is the actual summary text)
                for j in range(i + 1, min(i + 6, len(paras))):
                    if paras[j].text.strip():
                        replace_para_text(paras[j], summary)
                        summary_replaced = True
                        break
                break

        # Strategy B: paragraph that STARTS with "Summary" then has the content inline
        # e.g. "Summary\n\nSenior DevOps..." — replace it with heading + new summary
        if not summary_replaced:
            for i, para in enumerate(paras[:40]):
                text_lower = para.text.lower().strip()
                starts_with_summary = any(
                    text_lower.startswith(kw + "\n") or text_lower.startswith(kw + " ")
                    for kw in summary_keywords
                )
                if starts_with_summary and len(para.text.strip()) > 80:
                    # This paragraph IS the summary (heading merged with content)
                    replace_para_text(para, summary)
                    summary_replaced = True
                    break

        # Strategy C: fallback — find the longest paragraph in first 40 paras
        if not summary_replaced:
            best_i, best_len = -1, 0
            for i, para in enumerate(paras[:40]):
                ln = len(para.text.strip())
                if ln > best_len:
                    best_len = ln
                    best_i = i
            if best_i >= 0 and best_len > 100:
                replace_para_text(paras[best_i], summary)
                summary_replaced = True

    # ── Professional skills bullets replacement ──────────────────────────────
    if skills_bullets:
        skills_keywords = ["professional skills", "technical skills", "core skills",
                           "skills", "competencies", "expertise", "key skills"]
        # Stop scanning if we hit one of these section boundaries
        section_stopwords = ["education", "certif", "award", "experience",
                             "employment", "career", "project"]

        for i, para in enumerate(paras):
            text_lower = para.text.lower().strip()
            if any(kw == text_lower for kw in skills_keywords):
                # Collect ALL candidate bullet paragraphs after the heading
                # until we hit a real section boundary (heading / section title)
                candidate_indices = []
                consecutive_non_bullets = 0
                for j in range(i + 1, min(i + 50, len(paras))):
                    p = paras[j]
                    pt = p.text.strip()
                    pt_lower = pt.lower()

                    # Hard stop: section heading or section boundary word
                    is_section_boundary = (
                        p.style.name.lower().startswith("heading") or
                        any(sw in pt_lower for sw in section_stopwords)
                    )
                    if is_section_boundary and j > i + 2:
                        break

                    # Is it a skills-type line?
                    is_skills_line = (
                        pt.startswith("•") or
                        pt.startswith("-") or
                        pt.startswith("·") or
                        ":" in pt[:50] or       # "Cloud Platforms: ..."
                        p.style.name.lower().startswith("list") or
                        (len(pt) > 30 and len(pt) < 300)  # typical skills bullet length
                    )
                    if is_skills_line and pt:
                        candidate_indices.append(j)
                        consecutive_non_bullets = 0
                    elif not pt:
                        pass  # allow blank lines
                    else:
                        consecutive_non_bullets += 1
                        if consecutive_non_bullets > 2:
                            break  # too many non-bullet lines in a row

                # Replace up to len(skills_bullets) candidates
                for k, new_bullet in enumerate(skills_bullets):
                    if k < len(candidate_indices):
                        replace_para_text(paras[candidate_indices[k]], new_bullet)
                break

    # ── Job bullets replacement ──────────────────────────────────────────────
    if job_bullets:
        for company_name, new_bullets in job_bullets.items():
            company_lower = company_name.lower().strip()
            # Find the company heading paragraph
            company_para_idx = None
            for i, text in enumerate(para_texts):
                if company_lower in text.lower() and len(text) < 120:
                    company_para_idx = i
                    break

            if company_para_idx is None:
                continue

            # Find bullet paragraphs after the company heading
            # Look within next 30 paragraphs for bullet lines
            bullet_indices = []
            for i in range(company_para_idx + 1, min(company_para_idx + 35, len(paras))):
                p = paras[i]
                pt = p.text.strip()
                is_bullet = (
                    pt.startswith("•") or
                    pt.startswith("·") or
                    pt.startswith("-") or
                    p.style.name.lower().startswith("list") or
                    (len(pt) > 40 and not any(
                        kw in pt.lower() for kw in
                        ["ltd", "inc", "plc", "corp", "limited", "jan", "feb",
                         "mar", "apr", "may", "jun", "jul", "aug", "sep",
                         "oct", "nov", "dec", "20", "19"]
                    ) and ":" not in pt[:30])
                )
                # Stop if we hit another company/section heading
                is_heading = (
                    len(pt) < 60 and
                    any(kw in pt.lower() for kw in
                        ["education", "certification", "award", "volunteer"]) or
                    p.style.name.lower().startswith("heading")
                )
                if is_heading and i > company_para_idx + 2:
                    break
                if is_bullet and pt:
                    bullet_indices.append(i)

            # Replace up to min(available_bullets, new_bullets) paragraphs
            for k, new_bullet in enumerate(new_bullets):
                if k < len(bullet_indices):
                    replace_para_text(paras[bullet_indices[k]], new_bullet)

    # ── Save to bytes ────────────────────────────────────────────────────────
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ─── Extract text from DOCX bytes ────────────────────────────────────────────

def extract_text_from_docx_bytes(cv_bytes: bytes) -> str:
    """Extract plain text from DOCX bytes for scoring."""
    doc = Document(io.BytesIO(cv_bytes))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    return "\n".join(lines)
